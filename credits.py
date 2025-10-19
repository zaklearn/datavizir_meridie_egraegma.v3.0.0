"""
Datavizir Analytics - Educational Assessment Platform
Copyright (c) 2024-2025 Zakaria Benhoumad & Meridie ONG
Licensed under MIT License with Attribution Requirement

This module handles credits, licensing, and attribution display.
"""

import streamlit as st
from datetime import datetime
import hashlib


# ==================== CONFIGURATION ====================
CREDITS_CONFIG = {
    "author": "Zakaria Benhoumad",
    "website": "bendatainsights.cloud",
    "organization": "ONG Meridie",
    "org_website": "meridie.org",
    "project_name": "Datavizir Analytics",
    "version": "2.2.2",
    "license": "MIT License with Attribution",
    "year": "2024-2025",
    "github": "https://github.com/zakaria-benhoumad/datavizir-analytics",
    "contact": "contact@bendatainsights.cloud"
}


# ==================== HASH DE VÉRIFICATION ====================
def generate_app_hash():
    """Génère un hash unique pour vérifier l'intégrité de l'application."""
    content = f"{CREDITS_CONFIG['author']}{CREDITS_CONFIG['organization']}{CREDITS_CONFIG['version']}"
    return hashlib.sha256(content.encode()).hexdigest()[:16]


APP_HASH = generate_app_hash()


# ==================== FONCTIONS D'AFFICHAGE ====================
def get_translations(language="en"):
    """
    Récupère les traductions depuis config.py.
    Fallback vers l'anglais si la langue n'est pas disponible.
    """
    try:
        from config import translations
        return translations.get(language, translations.get("en", {}))
    except ImportError:
        # Fallback basique si config.py n'est pas disponible
        return {
            "credits_designed_by": "Designed by" if language == "en" else "Conçu par",
            "credits_managed_by": "Managed by" if language == "en" else "Géré par",
            "credits_open_source": "Open Source",
            "credits_development": "Development" if language == "en" else "Développement",
            "credits_organization": "Organization" if language == "en" else "Organisation",
            "credits_version": "Version",
            "credits_license": "License" if language == "en" else "Licence",
        }


def show_credits_sidebar(language="en"):
    """
    Affiche un bloc de crédits compact dans la sidebar.
    
    Args:
        language (str): Code langue (en/fr/ar)
    """
    t = get_translations(language)
    
    st.sidebar.markdown("---")
    st.sidebar.markdown(f"""
    <div style='text-align: center; padding: 10px; background-color: rgba(0,0,0,0.05); border-radius: 5px;'>
        <p style='margin: 0; font-size: 0.8em; color: #666;'>
            <strong> {CREDITS_CONFIG['project_name']}</strong><br>
            v{CREDITS_CONFIG['version']}<br>
            <br>
             <strong>{t.get('credits_designed_by', 'Designed by')}</strong><br>
            <a href='https://{CREDITS_CONFIG["website"]}' target='_blank' style='color: #1f77b4; text-decoration: none;'>
                {CREDITS_CONFIG['author']}
            </a><br>
             {CREDITS_CONFIG['website']}<br>
            <br>
             <strong>{t.get('credits_managed_by', 'Managed by')}</strong><br>
            <a href='https://{CREDITS_CONFIG["org_website"]}' target='_blank' style='color: #1f77b4; text-decoration: none;'>
                {CREDITS_CONFIG['organization']}
            </a><br>
             {CREDITS_CONFIG['org_website']}<br>
            <br>
            <span style='font-size: 0.7em;'>
                 {t.get('credits_open_source', 'Open Source')}<br>
                © {CREDITS_CONFIG['year']}
            </span>
        </p>
    </div>
    """, unsafe_allow_html=True)


def show_credits_footer(language="en"):
    """
    Affiche un footer complet en bas de page.
    
    Args:
        language (str): Code langue (en/fr/ar)
    """
    t = get_translations(language)
    
    st.markdown("---")
    
    col1, col2, col3 = st.columns([1, 2, 1])
    
    with col1:
        st.markdown(f"""
        {t.get('credits_development', 'Development')}    
        {CREDITS_CONFIG['author']}  
        [{CREDITS_CONFIG['website']}](https://{CREDITS_CONFIG['website']})
        """)
    
    with col2:
        platform_desc = t.get('credits_platform_description', 'Educational Assessment Analytics Platform')
        st.markdown(f"""
        <div style='text-align: center;'>
            <strong> {CREDITS_CONFIG['project_name']}</strong><br>
            {platform_desc}<br>
            <span style='font-size: 0.9em; color: #666;'>
                {t.get('credits_version', 'Version')} {CREDITS_CONFIG['version']} | 
                {t.get('credits_license', 'License')}: MIT<br>
                © {CREDITS_CONFIG['year']} - {t.get('credits_all_rights', 'All rights reserved')}
            </span>
        </div>
        """, unsafe_allow_html=True)
    
    with col3:
        st.markdown(f"""
        {t.get('credits_organization', 'Organization')}    
        {CREDITS_CONFIG['organization']}  
        [{CREDITS_CONFIG['org_website']}](https://{CREDITS_CONFIG['org_website']})
        """)
    
    # Hash de vérification (discret)
    st.markdown(
        f"<p style='text-align: center; font-size: 0.7em; color: #999; margin-top: 10px;'>App ID: {APP_HASH}</p>",
        unsafe_allow_html=True
    )


def show_credits_full_page(language="en"):
    """
    Affiche une page complète de crédits et informations légales.
    
    Args:
        language (str): Code langue (en/fr/ar)
    """
    t = get_translations(language)
    
    st.title(f" {t.get('credits_page_title', 'Credits and Legal Information')}")
    
    # Section Développement
    st.header(f" {t.get('credits_conception_dev', 'Conception and Development')}")
    
    dev_desc = t.get('credits_dev_description', f"""
This application was designed and developed by **{CREDITS_CONFIG['author']}**.

-  Website: [{CREDITS_CONFIG['website']}](https://{CREDITS_CONFIG['website']})
-  Contact: {CREDITS_CONFIG['contact']}
-  Services: Data Analytics, Business Intelligence, Educational Technology

**Ben Data Insights** specializes in creating data analysis solutions 
for the education sector and social impact organizations.
    """)
    
    st.markdown(dev_desc.format(
        author=CREDITS_CONFIG['author'],
        website=CREDITS_CONFIG['website'],
        contact=CREDITS_CONFIG['contact']
    ))
    
    # Section Organisation
    st.header(f" {t.get('credits_management_deployment', 'Management and Deployment')}")
    
    org_desc = t.get('credits_org_description', f"""
This platform is managed and deployed by **{CREDITS_CONFIG['organization']}**.

- Website: [{CREDITS_CONFIG['org_website']}](https://{CREDITS_CONFIG['org_website']})
- Mission: Improve education through data analysis
- Impact: Support educational systems in developing countries
    """)
    
    st.markdown(org_desc.format(
        organization=CREDITS_CONFIG['organization'],
        org_website=CREDITS_CONFIG['org_website']
    ))
    
    # Section Licence
    st.header(f" {t.get('credits_license_section', 'Open Source License')}")
    
    license_text = t.get('credits_license_intro', 
        f"**{CREDITS_CONFIG['project_name']}** is distributed under **MIT License with Attribution**.")
    st.markdown(license_text.format(
        project_name=CREDITS_CONFIG['project_name'],
        license=CREDITS_CONFIG['license']
    ))
    
    # Section Citation
    st.header(f" {t.get('credits_academic_citation', 'Academic Citation')}")
    
    citation = t.get('credits_citation_format', f"""
**To cite this application in your academic work:**

Benhoumad, Z. ({datetime.now().year}). *{CREDITS_CONFIG['project_name']}: Educational Assessment 
Analytics Platform EGRA/EGMA* (Version {CREDITS_CONFIG['version']}) [Software]. 
ONG Meridie. https://{CREDITS_CONFIG['org_website']}
    """)
    
    st.info(citation.format(
        year=datetime.now().year,
        project_name=CREDITS_CONFIG['project_name'],
        version=CREDITS_CONFIG['version'],
        org_website=CREDITS_CONFIG['org_website']
    ))
    
    # Section Contact
    st.header(f" {t.get('credits_contact_support', 'Contact and Support')}")
    
    contact_info = t.get('credits_contact_info', f"""
For questions, suggestions or collaboration:

-  Email: {CREDITS_CONFIG['contact']}
-  Web: [{CREDITS_CONFIG['website']}](https://{CREDITS_CONFIG['website']})
-  LinkedIn: Zakaria Benhoumad
-  GitHub: [Datavizir Analytics]({CREDITS_CONFIG['github']})
    """)
    
    st.markdown(contact_info.format(
        contact=CREDITS_CONFIG['contact'],
        website=CREDITS_CONFIG['website'],
        github=CREDITS_CONFIG['github']
    ))
    
    # Footer avec hash
    st.markdown("---")
    st.markdown(
        f"<p style='text-align: center; color: #666;'>© {CREDITS_CONFIG['year']} {CREDITS_CONFIG['author']} "
        f"& {CREDITS_CONFIG['organization']} | {t.get('credits_version', 'Version')} {CREDITS_CONFIG['version']} | App ID: {APP_HASH}</p>",
        unsafe_allow_html=True
    )


def add_credits_to_word_report(doc, language="en"):
    """
    Ajoute les crédits à un document Word.
    
    Args:
        doc: Document python-docx
        language (str): Code langue (en/fr/ar)
    
    Returns:
        doc: Document avec crédits ajoutés
    """
    try:
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        st.warning("python-docx non disponible. Les crédits ne seront pas ajoutés au rapport Word.")
        return doc
    
    t = get_translations(language)
    
    # Ajouter un saut de page
    doc.add_page_break()
    
    # Titre
    heading = doc.add_heading(t.get('credits_word_title', 'Credits and Information'), level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Développement
    doc.add_heading(t.get('credits_conception_dev', 'Conception and Development'), level=2)
    p1 = doc.add_paragraph()
    p1.add_run(t.get('credits_developed_by', 'Developed by') + ": ").bold = True
    p1.add_run(f"{CREDITS_CONFIG['author']}\n")
    p1.add_run(t.get('credits_website', 'Website') + f": {CREDITS_CONFIG['website']}\n")
    p1.add_run(t.get('credits_email', 'Email') + f": {CREDITS_CONFIG['contact']}\n")
    
    # Organisation
    doc.add_heading(t.get('credits_management_deployment', 'Management and Deployment'), level=2)
    p2 = doc.add_paragraph()
    p2.add_run(t.get('credits_managed_by', 'Managed by') + ": ").bold = True
    p2.add_run(f"{CREDITS_CONFIG['organization']}\n")
    p2.add_run(t.get('credits_website', 'Website') + f": {CREDITS_CONFIG['org_website']}\n")
    
    # Licence
    doc.add_heading(t.get('credits_license', 'License'), level=2)
    p3 = doc.add_paragraph()
    p3.add_run(f"{CREDITS_CONFIG['project_name']}\n")
    p3.add_run(t.get('credits_version', 'Version') + f" {CREDITS_CONFIG['version']}\n")
    p3.add_run(f"{CREDITS_CONFIG['license']}\n")
    p3.add_run(f"© {CREDITS_CONFIG['year']} - " + t.get('credits_all_rights', 'All rights reserved') + "\n")
    
    # Note d'attribution
    doc.add_paragraph()
    note = doc.add_paragraph()
    note.add_run(t.get('credits_note', 'Note') + ": ").bold = True
    note.add_run(t.get('credits_attribution_note', 
        "This report was generated with Datavizir Analytics. "
        "Use, modification or distribution of this software "
        "requires maintaining this attribution in accordance with the MIT license."))
    note_format = note.paragraph_format
    note_format.space_before = Pt(12)
    note_format.space_after = Pt(12)
    
    # Footer discret
    footer = doc.add_paragraph()
    footer_run = footer.add_run(f"\nApp ID: {APP_HASH} | {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(150, 150, 150)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    return doc


def verify_app_integrity():
    """Vérifie l'intégrité de l'application."""
    current_hash = generate_app_hash()
    return current_hash == APP_HASH


def show_license_warning_if_modified(language="en"):
    """
    Affiche un avertissement si les crédits ont été supprimés ou modifiés.
    
    Args:
        language (str): Code langue (en/fr/ar)
    """
    t = get_translations(language)
    
    if not verify_app_integrity():
        st.sidebar.warning(t.get('credits_attribution_warning', """
⚠️ **Attribution Required**

This software is open source but requires attribution to original authors.
        """))


# ==================== INTÉGRATION RAPIDE ====================
def initialize_credits(location="sidebar", language="en"):
    """
    Fonction d'initialisation rapide des crédits.
    
    Args:
        location (str): "sidebar", "footer", ou "both"
        language (str): Code langue (en/fr/ar)
    """
    show_license_warning_if_modified(language)
    
    if location in ["sidebar", "both"]:
        show_credits_sidebar(language)
    
    if location in ["footer", "both"]:
        show_credits_footer(language)




def show_credits_fixed_footer(language="en"):
  
    """
    Affiche un footer fixe en bas de page en réutilisant show_credits_footer.
    Ajoute uniquement le CSS pour le rendre fixe.
    
    Args:
        language (str): Code langue (en/fr/ar)
    """
    # Conteneur pour le footer avec un ID unique
    st.markdown('<div id="footer-container">', unsafe_allow_html=True)
    
    # Appeler la fonction footer existante
    show_credits_footer(language)
    
    st.markdown('</div>', unsafe_allow_html=True)
    
    # CSS pour rendre le footer fixe
    st.markdown("""
        <style>
        /* Rendre le footer fixe en bas */
        #footer-container {
            position: fixed;
            left: 0;
            bottom: 0;
            width: 100%;
            background-color: #0E1117;
            padding: 10px 20px;
            z-index: 999;
            border-top: 2px solid #262730;
            box-shadow: 0 -2px 10px rgba(0,0,0,0.3);
        }
        
        /* Ajuster les liens dans le footer */
        #footer-container a {
            color: #FF4B4B !important;
            text-decoration: none;
        }
        
        #footer-container a:hover {
            text-decoration: underline;
            color: #FF6B6B !important;
        }
        
        /* Espacement pour éviter que le footer cache le contenu */
        .main .block-container {
            padding-bottom: 200px !important;
        }
        
        /* Ajuster le style des colonnes dans le footer */
        #footer-container [data-testid="column"] {
            background-color: transparent !important;
        }
        </style>
    """, unsafe_allow_html=True)