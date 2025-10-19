import streamlit as st
import pandas as pd
import plotly.express as px
from docx import Document
from docx.shared import Inches, RGBColor
import tempfile
import os
import warnings
from io import BytesIO
from dotenv import load_dotenv
import google.generativeai as genai

# Supprimer les warnings FutureWarning de pandas
warnings.filterwarnings('ignore', category=FutureWarning)

# Import configuration depuis le fichier de config principal
from config import translations

# Import configuration depuis le fichier de config principal
from config import translations

# Import du module de crédits
try:
    from credits import initialize_credits, add_credits_to_word_report
    CREDITS_AVAILABLE = True
except ImportError:
    CREDITS_AVAILABLE = False
    st.warning("⚠️ Module 'credits.py' non trouvé. Les crédits ne seront pas affichés.")


def dataframe_to_markdown(df: pd.DataFrame) -> str:
    """
    Convertit un DataFrame en format Markdown sans dépendance externe.
    
    Args:
        df (pd.DataFrame): DataFrame à convertir
        
    Returns:
        str: Représentation Markdown du DataFrame
    """
    # Créer les en-têtes
    headers = df.columns.tolist()
    markdown = "| " + " | ".join(str(h) for h in headers) + " |\n"
    markdown += "|" + "|".join(["---" for _ in headers]) + "|\n"
    
    # Ajouter les lignes
    for _, row in df.iterrows():
        markdown += "| " + " | ".join(str(v) for v in row.values) + " |\n"
    
    return markdown


def generate_gemini_interpretation(df_zero_scores: pd.DataFrame, t: dict) -> str:
    """
    Génère une interprétation éducative et des recommandations en utilisant l'API Gemini.
    VERSION CORRIGÉE - Multilingue
    
    Args:
        df_zero_scores (pd.DataFrame): DataFrame avec l'analyse des scores nuls
        t (dict): Dictionnaire de traductions
        
    Returns:
        str: Le texte de l'interprétation générée par l'IA, ou None si erreur
    """
    import time
    
    # Vérifier si l'API Gemini est disponible
    if not GEMINI_AVAILABLE:
        st.error(t.get("api_not_configured", "❌ **API Gemini non configurée**"))
        st.info(t.get("api_activation_steps", """
**Pour activer l'interprétation par IA, suivez ces étapes :**

1. Créez un fichier `.env` à la racine de votre projet
2. Ajoutez votre clé API Gemini dans ce fichier :
   ```
   GEMINI_API_KEY=votre_clé_api_ici
   ```
3. Obtenez une clé API gratuite sur [Google AI Studio](https://aistudio.google.com)
4. Redémarrez l'application

**En attendant**, vous pouvez consulter les tableaux et graphiques ci-dessus qui fournissent déjà des informations détaillées sur les performances.
        """))
        return None
    
    # Convertir le DataFrame en Markdown
    df_for_markdown = df_zero_scores[[
        t.get("task_column", "Task"),
        t.get("count_column", "Count of Zeros"),
        t.get("percentage_column", "Percentage of Zero Scores")
    ]].copy()
    
    data_as_markdown = dataframe_to_markdown(df_for_markdown)
    
    # ✅ UTILISER LE TEMPLATE DE PROMPT SELON LA LANGUE
    prompt_template_base = t.get("gemini_prompt_template", """**Context:** You are an internationally renowned expert in educational sciences, specialized in the analysis of large-scale assessments such as EGRA. Your analysis must be rigorous, evidence-based, and your recommendations must be practical for teachers.

**Raw Data to Analyze:** The table below shows the percentage of students who obtained a zero score for several fundamental assessment tasks. A zero score represents a complete absence of the measured skill.

```markdown
{data_as_markdown}
```

**Your Mission:** Write a comprehensive diagnostic analysis report in English. Your response must be structured in three distinct sections in Markdown format.

## 1. Pedagogical Interpretation

**Summary:** Begin with a 2-3 sentence synthesis of the general state of skills, identifying critical areas of strength and weakness.

**Concerning Areas:** Identify the most alarming skills (those with the highest percentages). Explain in detail why these deficits are critical for the student's future development. Create causal links between skills.

**Stability Points:** Briefly mention the skills that seem acquired (those with the lowest percentages).

## 2. Actionable Recommendations

**Priority Recommendations:** Propose very concrete and targeted intervention strategies for the weakest skills.

**Implementation Strategies:** Provide advice on how to integrate these recommendations (differentiation, small groups, etc.).

**Assessment Recommendations:** Suggest a follow-up plan to measure progress.

## 3. Reliable Sources and References

To give credibility to your analysis, cite at least two recognized academic or institutional sources that support your recommendations. List them clearly at the end.""")
    
    # Formater le prompt avec les données
    prompt_template = prompt_template_base.format(data_as_markdown=data_as_markdown)
    
    # Configuration du retry
    max_retries = 3
    retry_delay = 20  # secondes
    
    for attempt in range(max_retries):
        try:
            if attempt > 0:
                st.info(t.get("retry_message", "⏳ Nouvelle tentative ({attempt}/{max_retries}) dans {delay} secondes...").format(
                    attempt=attempt + 1,
                    max_retries=max_retries,
                    delay=retry_delay
                ))
                time.sleep(retry_delay)
            
            with st.spinner(t.get("generating_interpretation", "🤖 L'IA analyse les résultats...")):
                model = genai.GenerativeModel('gemini-2.5-pro')
                response = model.generate_content(prompt_template)
                
                # Afficher l'interprétation
                st.subheader(t.get("interpretation_title", "📝 Educational Interpretation"))
                st.markdown(response.text)
                
                # Retourner le texte pour l'utiliser dans le rapport Word
                return response.text
                
        except Exception as e:
            error_message = str(e)
            
            # Détecter spécifiquement l'erreur 429 (quota dépassé)
            if "429" in error_message or "quota" in error_message.lower():
                if attempt < max_retries - 1:
                    st.warning(t.get("quota_retry", "⚠️ Limite de quota API atteinte. Nouvelle tentative automatique dans {delay} secondes...").format(
                        delay=retry_delay
                    ))
                    retry_delay *= 2  # Backoff exponentiel
                else:
                    st.error(t.get("quota_exceeded", "❌ **Quota API Gemini dépassé**"))
                    st.info(t.get("quota_solutions", """
**Solutions possibles :**
1. 🕐 Attendez quelques minutes avant de réessayer
2. 🔑 Vérifiez votre plan API Gemini sur [Google AI Studio](https://aistudio.google.com)
3. 💳 Considérez passer à un plan payant pour des quotas plus élevés
4. 📊 Pour le moment, vous pouvez consulter les tableaux et graphiques ci-dessus

**Limites du niveau gratuit :**
- 2 requêtes par minute
- 1 500 requêtes par jour

[En savoir plus sur les quotas](https://ai.google.dev/gemini-api/docs/rate-limits)
                    """))
                    return None
            else:
                # Autre type d'erreur
                st.error(t.get("error_generating_report", "❌ Erreur lors de l'appel à l'API Gemini: {error}").format(
                    error=error_message
                ))
                if attempt < max_retries - 1:
                    st.warning(t.get("retry_message", "⏳ Nouvelle tentative dans {delay} secondes...").format(
                        attempt=attempt + 1,
                        max_retries=max_retries,
                        delay=retry_delay
                    ))
                else:
                    st.info(t.get("verification_suggestions", """
**Vérifications suggérées :**
- ✅ Votre clé API est correcte dans le fichier `.env`
- ✅ Vous avez une connexion internet active
- ✅ L'API Gemini est accessible depuis votre région
                    """))
                    return None
    
    return None


def create_complete_word_report(df_zero_scores: pd.DataFrame, fig, t: dict, ai_interpretation: str = None, language: str = "en") -> Document:
    """
    Crée un rapport Word COMPLET avec tableaux, graphiques ET interprétation IA.
    VERSION CORRIGÉE - Multilingue
    
    Args:
        df_zero_scores (pd.DataFrame): DataFrame avec l'analyse des scores nuls
        fig: Figure Plotly du graphique
        t (dict): Dictionnaire de traductions
        ai_interpretation (str): Texte de l'interprétation IA (optionnel)
        language (str): Code de langue (en/fr/ar)
        
    Returns:
        Document: Document Word complet
    """
    from datetime import datetime
    
    doc = Document()
    
    # ========== PAGE DE GARDE ==========
    doc.add_heading(t.get("title_zero_scores", "Zero Scores Analysis"), level=1)
    
    # Date et heure du rapport
    date_str = datetime.now().strftime("%d/%m/%Y %H:%M")
    doc.add_paragraph(f"📅 {t.get('report_date', 'Report date')}: {date_str}")
    doc.add_paragraph("_" * 50)
    doc.add_paragraph()
    
    # ========== RÉSUMÉ EXÉCUTIF ==========
    doc.add_heading(t.get("executive_summary", "Executive Summary"), level=2)
    
    # Calculer les statistiques globales
    total_tasks = len(df_zero_scores)
    percentage_col = t.get("percentage_column", "Percentage of Zero Scores")
    avg_percentage = df_zero_scores[percentage_col].mean()
    critical_count = len(df_zero_scores[df_zero_scores[percentage_col] >= 30])
    concerning_count = len(df_zero_scores[
        (df_zero_scores[percentage_col] >= 20) &
        (df_zero_scores[percentage_col] < 30)
    ])
    
    # ✅ UTILISER LES TRADUCTIONS AU LIEU DE TEXTE CODÉ EN DUR
    summary_text = f"""
{t.get("report_intro_text", "This report analyzes zero scores obtained by students on {total_tasks} assessment tasks.").format(total_tasks=total_tasks)}

{t.get("key_stats_title", "Key Statistics:")}
• {t.get("avg_zero_percentage", "Average percentage of zero scores: {avg_percentage:.1f}%").format(avg_percentage=avg_percentage)}
• {t.get("critical_tasks_count", "Critical tasks (>30% zeros): {critical_count}").format(critical_count=critical_count)}
• {t.get("concerning_tasks_count", "Concerning tasks (20-30% zeros): {concerning_count}").format(concerning_count=concerning_count)}

{t.get("zero_score_meaning", "A zero score indicates a complete absence of mastery of the assessed skill and requires particular attention.")}
    """
    doc.add_paragraph(summary_text)
    doc.add_paragraph()
    
    # ========== TABLEAU DES RÉSULTATS ==========
    doc.add_heading(t.get("table_zero_scores", "Proportion of Students with Zero Scores"), level=2)
    
    # Créer le tableau
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    
    # En-têtes
    header_cells = table.rows[0].cells
    header_cells[0].text = t.get("task_column", "Task")
    header_cells[1].text = t.get("count_column", "Count of Zeros")
    header_cells[2].text = t.get("percentage_column", "Percentage of Zero Scores")
    
    # Rendre les en-têtes en gras
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Ajouter les données
    task_col = t.get("task_column", "Task")
    count_col = t.get("count_column", "Count of Zeros")
    
    for _, row in df_zero_scores.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row[task_col])
        row_cells[1].text = str(row[count_col])
        percentage_val = row[percentage_col]
        row_cells[2].text = f"{percentage_val}%"
        
        # Colorer selon le seuil en utilisant RGBColor
        if percentage_val >= 30:
            # Rouge
            row_cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(220, 20, 60)
        elif percentage_val >= 20:
            # Orange
            row_cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 140, 0)
    
    doc.add_paragraph()
    
    # ========== GRAPHIQUE ==========
    doc.add_heading(t.get("visualization_title", "Visualization"), level=2)
    
    # Exporter le graphique en image
    with tempfile.TemporaryDirectory() as tmp_dir:
        img_path = os.path.join(tmp_dir, "zero_scores_chart.png")
        try:
            fig.write_image(img_path, width=1200, height=600)
            doc.add_picture(img_path, width=Inches(6.5))
        except Exception as e:
            doc.add_paragraph(f"[{t.get('visualization_title', 'Visualization')} not available: {str(e)}]")
    
    doc.add_paragraph()
    
    # ========== INTERPRÉTATION IA (si disponible) ==========
    if ai_interpretation:
        doc.add_heading(t.get("interpretation_title", "Educational Interpretation"), level=2)
        doc.add_paragraph(t.get("ai_interpretation_notice", "🤖 This interpretation was generated by artificial intelligence (Gemini)"))
        doc.add_paragraph("_" * 50)
        
        # Ajouter l'interprétation IA (en préservant le formatage Markdown basique)
        for line in ai_interpretation.split('\n'):
            if line.strip().startswith('##'):
                # Titre de niveau 2
                doc.add_heading(line.replace('##', '').strip(), level=3)
            elif line.strip().startswith('**') and line.strip().endswith('**'):
                # Texte en gras
                p = doc.add_paragraph()
                p.add_run(line.strip().replace('**', '')).bold = True
            elif line.strip().startswith('- ') or line.strip().startswith('* '):
                # Liste à puces
                doc.add_paragraph(line.strip()[2:], style='List Bullet')
            elif line.strip():
                # Paragraphe normal
                doc.add_paragraph(line.strip())
        
        doc.add_paragraph()
    else:
        # ========== RECOMMANDATIONS DE BASE (si pas d'IA) ==========
        doc.add_heading(t.get("recommendations_title", "Recommendations"), level=2)
        
        # Catégoriser les tâches
        critical_tasks = df_zero_scores[df_zero_scores[percentage_col] >= 30]
        concerning_tasks = df_zero_scores[
            (df_zero_scores[percentage_col] >= 20) &
            (df_zero_scores[percentage_col] < 30)
        ]
        
        if len(critical_tasks) > 0:
            doc.add_heading(t.get("critical_areas_title", "🔴 Critical Areas (>30% zero scores)"), level=3)
            doc.add_paragraph(t.get("critical_areas_description", 
                "These skills require immediate intervention with intensive and targeted teaching programs."))
            for _, task in critical_tasks.iterrows():
                doc.add_paragraph(
                    f"• {task[task_col]}: {task[percentage_col]}%",
                    style='List Bullet'
                )
        
        if len(concerning_tasks) > 0:
            doc.add_heading(t.get("concerning_areas_title", "🟠 Concerning Areas (20-30% zero scores)"), level=3)
            doc.add_paragraph(t.get("concerning_areas_description",
                "These skills require significant reinforcement within regular instruction."))
            for _, task in concerning_tasks.iterrows():
                doc.add_paragraph(
                    f"• {task[task_col]}: {task[percentage_col]}%",
                    style='List Bullet'
                )
        
        doc.add_paragraph()
        doc.add_heading(t.get("general_strategies_title", "General Intervention Strategies"), level=3)
        doc.add_paragraph(t.get("strategy_1", "Differentiated instruction in small groups"), style='List Number')
        doc.add_paragraph(t.get("strategy_2", "In-depth diagnostic assessment to identify specific gaps"), style='List Number')
        doc.add_paragraph(t.get("strategy_3", "Early and intensive intervention for struggling students"), style='List Number')
        doc.add_paragraph(t.get("strategy_4", "Regular progress monitoring (every 2-3 weeks)"), style='List Number')
        doc.add_paragraph(t.get("strategy_5", "Collaboration with families for home support"), style='List Number')
    
    # ========== NOTES MÉTHODOLOGIQUES ==========
    doc.add_page_break()
    doc.add_heading(t.get("methodology_title", "Methodological Notes"), level=2)
    
    # ✅ UTILISER LES TRADUCTIONS
    methodology_text = f"""
{t.get("methodology_intro", "This report analyzes zero scores in EGRA/EGMA assessments according to the following criteria:")}

{t.get("interpretation_thresholds", "Interpretation Thresholds:")}
• {t.get("threshold_acceptable", "Acceptable: < 10% zero scores")}
• {t.get("threshold_monitor", "To monitor: 10-20% zero scores")}
• {t.get("threshold_concerning", "Concerning: 20-30% zero scores")}
• {t.get("threshold_critical", "Critical: > 30% zero scores")}

{t.get("methodology_explanation", "A high percentage of zero scores indicates that many students have not acquired the fundamental skills being assessed. These gaps can compromise future learning and require immediate attention.")}

{t.get("methodology_basis", "Recommendations are based on best practices in teaching primary reading and mathematics, as documented by educational research.")}
    """
    doc.add_paragraph(methodology_text)
    
    # ========== PIED DE PAGE ==========
    doc.add_paragraph()
    doc.add_paragraph("_" * 50)
    footer_text = f"{t.get('report_generated_by', 'Report generated by Datavizir Analytics')} - {date_str}"
    p = doc.add_paragraph(footer_text)
    p.alignment = 1  # Centre
    
    # ========== AJOUTER LES CRÉDITS ==========
    if CREDITS_AVAILABLE:
        doc = add_credits_to_word_report(doc)
    
    return doc

# ==================================================
# File: analyse2.py (SECTION MODIFIÉE)
# Intégration de gemini_config.py
# ==================================================

import streamlit as st
import pandas as pd
import plotly.express as px
from docx import Document
from docx.shared import Inches, RGBColor
import tempfile
import os
import warnings
from io import BytesIO

# ✅ SUPPRIMER CES LIGNES
# from dotenv import load_dotenv
# import google.generativeai as genai
# load_dotenv()
# api_key = os.getenv("GEMINI_API_KEY")
# GEMINI_AVAILABLE = False

# ✅ AJOUTER CETTE LIGNE
from gemini_config import get_gemini_config

# Supprimer les warnings FutureWarning de pandas
warnings.filterwarnings('ignore', category=FutureWarning)

# Import configuration depuis le fichier de config principal
from config import translations

# Import du module de crédits
try:
    from credits import initialize_credits, add_credits_to_word_report
    CREDITS_AVAILABLE = True
except ImportError:
    CREDITS_AVAILABLE = False
    st.warning("⚠️ Module 'credits.py' non trouvé. Les crédits ne seront pas affichés.")


# ==================================================
# FONCTION MODIFIÉE : generate_gemini_interpretation
# ==================================================

def generate_gemini_interpretation(df_zero_scores: pd.DataFrame, t: dict) -> str:
    """
    Génère une interprétation éducative en utilisant Gemini.
    VERSION REFACTORISÉE avec gemini_config.py
    
    Args:
        df_zero_scores (pd.DataFrame): DataFrame avec l'analyse des scores nuls
        t (dict): Dictionnaire de traductions
        
    Returns:
        str: Le texte de l'interprétation générée par l'IA, ou None si erreur
    """
    import time
    
    # ✅ RÉCUPÉRER LA CONFIG GEMINI
    gemini_config = get_gemini_config()
    model = gemini_config.get_model()
    
    # Vérifier si le modèle est disponible
    if not model:
        st.error(t.get("api_not_configured", "❌ **API Gemini non configurée**"))
        st.info(t.get("api_activation_steps", """
**Pour activer l'interprétation par IA :**
1. Configurez votre clé API dans la section "🤖 Gemini AI" ci-dessus
2. Obtenez une clé API gratuite sur [Google AI Studio](https://aistudio.google.com)
3. Validez la clé en cliquant sur "✅ Valider"

**En attendant**, vous pouvez consulter les tableaux et graphiques qui fournissent déjà des informations détaillées.
        """))
        return None
    
    # Convertir le DataFrame en Markdown
    df_for_markdown = df_zero_scores[[
        t.get("task_column", "Task"),
        t.get("count_column", "Count of Zeros"),
        t.get("percentage_column", "Percentage of Zero Scores")
    ]].copy()
    
    data_as_markdown = dataframe_to_markdown(df_for_markdown)
    
    # Utiliser le template de prompt selon la langue
    prompt_template_base = t.get("gemini_prompt_template", """**Context:** You are an internationally renowned expert in educational sciences, specialized in the analysis of large-scale assessments such as EGRA. Your analysis must be rigorous, evidence-based, and your recommendations must be practical for teachers.

**Raw Data to Analyze:** The table below shows the percentage of students who obtained a zero score for several fundamental assessment tasks. A zero score represents a complete absence of the measured skill.

```markdown
{data_as_markdown}
```

**Your Mission:** Write a comprehensive diagnostic analysis report. Your response must be structured in three distinct sections in Markdown format.

## 1. Pedagogical Interpretation

**Summary:** Begin with a 2-3 sentence synthesis of the general state of skills.

**Concerning Areas:** Identify the most alarming skills. Explain why these deficits are critical.

**Stability Points:** Mention the skills that seem acquired.

## 2. Actionable Recommendations

**Priority Recommendations:** Propose concrete intervention strategies.

**Implementation Strategies:** Provide advice on integration.

**Assessment Recommendations:** Suggest a follow-up plan.

## 3. Reliable Sources and References

Cite at least two recognized academic sources.""")
    
    prompt_template = prompt_template_base.format(data_as_markdown=data_as_markdown)
    
    # Configuration du retry
    max_retries = 3
    retry_delay = 20  # secondes
    
    for attempt in range(max_retries):
        try:
            if attempt > 0:
                st.info(t.get("retry_message", 
                    "⏳ Nouvelle tentative ({attempt}/{max_retries}) dans {delay} secondes...").format(
                    attempt=attempt + 1,
                    max_retries=max_retries,
                    delay=retry_delay
                ))
                time.sleep(retry_delay)
            
            with st.spinner(t.get("generating_interpretation", "🤖 L'IA analyse les résultats...")):
                # ✅ UTILISER LA MÉTHODE DE GÉNÉRATION AMÉLIORÉE
                response = model.generate_content(
                    prompt_template,
                    generation_config=gemini_config.get_generation_config()
                )
                
                # Afficher l'interprétation
                st.subheader(t.get("interpretation_title", "🔬 Educational Interpretation"))
                st.markdown(response.text)
                
                # Retourner le texte pour le rapport Word
                return response.text
                
        except Exception as e:
            error_message = str(e)
            
            # Détecter l'erreur 429 (quota dépassé)
            if "429" in error_message or "quota" in error_message.lower():
                if attempt < max_retries - 1:
                    st.warning(t.get("quota_retry", 
                        "⚠️ Limite de quota API atteinte. Nouvelle tentative dans {delay} secondes...").format(
                        delay=retry_delay
                    ))
                    retry_delay *= 2  # Backoff exponentiel
                else:
                    st.error(t.get("quota_exceeded", "❌ **Quota API Gemini dépassé**"))
                    st.info(t.get("quota_solutions", """
**Solutions possibles :**
1. 🕐 Attendez quelques minutes avant de réessayer
2. 🔑 Vérifiez votre plan API Gemini
3. 💳 Considérez passer à un plan payant
4. 📊 Consultez les tableaux et graphiques en attendant

**Limites du niveau gratuit :**
- 2 requêtes par minute
- 1 500 requêtes par jour
                    """))
                    return None
            else:
                # Autre type d'erreur
                st.error(t.get("error_generating_report", 
                    "❌ Erreur lors de l'appel à l'API Gemini: {error}").format(error=error_message))
                if attempt < max_retries - 1:
                    st.warning(t.get("retry_message", 
                        "⏳ Nouvelle tentative dans {delay} secondes...").format(
                        attempt=attempt + 1,
                        max_retries=max_retries,
                        delay=retry_delay
                    ))
                else:
                    st.info(t.get("verification_suggestions", """
**Vérifications suggérées :**
- ✅ Votre clé API est correcte
- ✅ Vous avez une connexion internet active
- ✅ L'API Gemini est accessible depuis votre région
                    """))
                    return None
    
    return None


# ==================================================
# FONCTION PRINCIPALE MODIFIÉE : show_zero_scores
# ==================================================

def show_zero_scores(df: pd.DataFrame, language: str) -> None:
    """
    Analyse et affiche la proportion de scores nuls.
    VERSION REFACTORISÉE avec nouvelle disposition UI
    
    Args:
        df (pd.DataFrame): Les données à analyser
        language (str): Langue sélectionnée (en/fr/ar/es)
    """
    t = translations[language]
    
    # Initialiser les crédits dans la sidebar
    if CREDITS_AVAILABLE:
        initialize_credits(location="sidebar", language=language)
    
    # Initialiser session_state
    if 'show_interpretation' not in st.session_state:
        st.session_state.show_interpretation = False
    if 'ai_interpretation_text' not in st.session_state:
        st.session_state.ai_interpretation_text = None
    if 'zero_scores_fig' not in st.session_state:
        st.session_state.zero_scores_fig = None
    
    # Définir les listes de colonnes EGRA et EGMA
    egra_columns = ["clpm", "phoneme", "sound_word", "cwpm", "listening", "orf", "comprehension"]
    egma_columns = ["number_id", "discrimin", "missing_number", "addition", "subtraction", "problems"]
    
    # Vérifier les colonnes disponibles
    available_columns = [col for col in egra_columns + egma_columns if col in df.columns]
    
    if not available_columns:
        st.error(t.get("no_assessment_columns", "No assessment columns found in the data."))
        return
    
    # Sélecteurs de variables
    st.subheader(t.get("select_variables", "📊 Select Variables"))
    
    # Multiselect pour EGRA
    available_egra = [col for col in egra_columns if col in df.columns]
    selected_egra = st.multiselect(
        t.get("egra_variables", "EGRA Variables:"),
        options=available_egra,
        default=available_egra[:3] if len(available_egra) > 3 else available_egra,
        format_func=lambda x: t["columns_of_interest"].get(x, x)
    )
    
    # Multiselect pour EGMA
    available_egma = [col for col in egma_columns if col in df.columns]
    selected_egma = st.multiselect(
        t.get("egma_variables", "EGMA Variables:"),
        options=available_egma,
        default=available_egma[:3] if len(available_egma) > 3 else available_egma,
        format_func=lambda x: t["columns_of_interest"].get(x, x)
    )
    
    # Combiner les variables sélectionnées
    selected_columns = selected_egra + selected_egma
    
    if selected_columns:
        try:
            # Calculer les scores nuls
            zero_scores = (df[selected_columns] == 0).sum()
            total_students = len(df)
            percentage_zero = ((zero_scores / total_students) * 100).round(2)
            
            # Créer un DataFrame pour l'affichage
            df_zero_scores = pd.DataFrame({
                "Task": [t.get("columns_of_interest", {}).get(col, col) for col in selected_columns],
                "Zero_Count": zero_scores.values,
                "Percentage": percentage_zero.values,
                "Task_Code": selected_columns
            })
            
            # Tableau des résultats
            st.subheader(t.get("table_zero_scores", "📋 Proportion of Students with Zero Scores"))
            styled_df = df_zero_scores.copy()
            styled_df.columns = [
                t.get("task_column", "Task"), 
                t.get("count_column", "Count of Zeros"),
                t.get("percentage_column", "Percentage of Zero Scores"),
                "Task_Code"
            ]
            st.dataframe(styled_df[styled_df.columns[:-1]], use_container_width=True)
            
            # Graphique de visualisation
            st.subheader(t.get("zero_scores_chart_title", "📊 Percentage of Students with Zero Scores by Task"))
            try:
                df_zero_scores_sorted = df_zero_scores.sort_values("Percentage", ascending=True).copy()
                df_zero_scores_sorted["Percentage"] = pd.to_numeric(df_zero_scores_sorted["Percentage"], errors='coerce')
                df_zero_scores_sorted["Task"] = df_zero_scores_sorted["Task"].astype(str)
                
                fig = px.bar(
                    df_zero_scores_sorted,
                    x="Percentage",
                    y="Task",
                    orientation="h",
                    text="Percentage",
                    color="Percentage",
                    color_continuous_scale="Viridis",
                    title=t.get("zero_scores_chart_title", "Percentage of Students with Zero Scores by Task")
                )
                
                fig.update_traces(texttemplate='%{text:.1f}%', textposition='outside')
                
                # Seuils critiques
                fig.add_vline(x=10, line_width=2, line_dash="dash", line_color="yellow", opacity=0.7)
                fig.add_vline(x=20, line_width=2, line_dash="dash", line_color="orange", opacity=0.7)
                fig.add_vline(x=30, line_width=2, line_dash="dash", line_color="red", opacity=0.7)
                
                # Annotations
                fig.add_annotation(x=10, y=0, text=t.get("acceptable_threshold", "Acceptable"), showarrow=False, yshift=-20, font=dict(size=10, color="yellow"))
                fig.add_annotation(x=20, y=0, text=t.get("concerning_threshold", "Concerning"), showarrow=False, yshift=-20, font=dict(size=10, color="orange"))
                fig.add_annotation(x=30, y=0, text=t.get("critical_threshold", "Critical"), showarrow=False, yshift=-20, font=dict(size=10, color="red"))
                
                fig.update_layout(
                    height=400,
                    xaxis_title=t.get("percentage_column", "Percentage of Zero Scores"),
                    yaxis_title=t.get("task_column", "Task"),
                    showlegend=False
                )
                
                st.session_state.zero_scores_fig = fig
                st.plotly_chart(fig, use_container_width=True)
            except Exception as e:
                st.error(f"Error creating visualization: {str(e)}")
            
            # ✅ NOUVELLE SECTION : CONFIGURATION GEMINI + ACTIONS
            st.divider()
            st.subheader(t.get("ai_analysis_section", "🤖 Analyse IA et Rapports"))
            
            # ✅ CONFIGURATION GEMINI (comme dans main.py mais en ligne)
            gemini_config = get_gemini_config()
            
            # Afficher la config dans un expander
            with st.expander("⚙️ " + t.get("gemini_configuration", "Configuration Gemini API"), expanded=not gemini_config.get_status().configured):
                gemini_config.render_inline_config_ui(t, language)
            
            # Status indicator
            status = gemini_config.get_status()
            if status.configured and status.model_loaded:
                st.success("✅ " + t.get("api_ready", "API Gemini configurée et opérationnelle"))
            elif status.api_key_set and not status.model_loaded:
                st.warning("⚠️ " + t.get("api_partial", "Clé présente mais modèle non chargé"))
            else:
                st.info("ℹ️ " + t.get("api_required", "Configuration requise pour l'analyse IA"))
            
            # Message d'aide
            st.info(f"""
            {t.get("usage_guide_title", "💡 Guide d'utilisation :")}
            - {t.get("usage_guide_ai", "**🔬 Interprétation IA** : Génère une analyse pédagogique détaillée avec recommandations")}
            - {t.get("usage_guide_report", "**📄 Rapport Complet** : Crée un document Word professionnel incluant tableaux, graphiques et interprétation IA")}
            """)
            
            # ✅ BOUTONS EN DISPOSITION VERTICALE
            # Bouton 1: Interprétation IA
            if status.configured:
                if st.button(
                    "🔬 " + t.get("generate_interpretation", "Générer l'Interprétation IA"),
                    type="primary",
                    use_container_width=True,
                    key="btn_interpretation"
                ):
                    st.session_state.show_interpretation = True
                    with st.spinner(t.get("generating_interpretation", "🤖 Génération de l'interprétation...")):
                        ai_text = generate_gemini_interpretation(styled_df, t)
                        st.session_state.ai_interpretation_text = ai_text
            else:
                st.button(
                    "🔬 " + t.get("generate_interpretation", "Générer l'Interprétation IA") + " 🔒",
                    disabled=True,
                    use_container_width=True,
                    help=t.get("api_locked_help", "Configurez d'abord votre clé API Gemini ci-dessus"),
                    key="btn_interpretation_locked"
                )
            
            # Bouton 2: Rapport Complet (juste en dessous)
            if st.button(
                "📄 " + t.get("export_complete_report", "Générer Rapport Complet (Word)"),
                use_container_width=True,
                key="btn_report"
            ):
                try:
                    if st.session_state.zero_scores_fig is not None:
                        with st.spinner(t.get("generating_report", "📝 Génération du rapport complet...")):
                            doc = create_complete_word_report(
                                styled_df,
                                st.session_state.zero_scores_fig,
                                t,
                                st.session_state.ai_interpretation_text,
                                language=language
                            )
                            
                            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                                doc.save(tmp.name)
                                with open(tmp.name, 'rb') as f:
                                    docx_data = f.read()
                                
                                st.download_button(
                                    "📥 " + t.get("download_complete_report", "Télécharger le rapport"),
                                    docx_data,
                                    "rapport_complet_zero_scores.docx",
                                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    key='download-complete-report',
                                    use_container_width=True
                                )
                        st.success(t.get("report_ready", "✅ Rapport généré avec succès!"))
                    else:
                        st.warning(t.get("wait_for_graph", "⚠️ Veuillez patienter, le graphique se charge..."))
                except Exception as e:
                    st.error(t.get("error_generating_report", "❌ Erreur: {error}").format(error=str(e)))
            
            # Section d'interprétation IA (affichée si générée)
            if st.session_state.show_interpretation and st.session_state.ai_interpretation_text:
                st.divider()
                # L'interprétation a déjà été affichée lors de la génération
                       
        except Exception as e:
            st.error(f"Error in zero scores analysis: {str(e)}")
    
    else:
        st.warning(t.get("warning_select_task", "Please select at least one task to analyze."))


# ... (garder les autres fonctions create_complete_word_report, dataframe_to_markdown inchangées)